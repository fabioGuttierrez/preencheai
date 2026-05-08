import logging
import re

from docx import Document

logger = logging.getLogger(__name__)


def gerar_contrato_docx(modelo_path: str, dados: dict, output_path: str) -> None:
    """
    Abre o modelo .docx, substitui todos os placeholders {{CHAVE}}
    pelos valores em `dados` e salva em output_path.

    Estratégia em dois passes:
    - Passo 1: substituição direta por run (preserva formatação bold/italic/cor).
    - Passo 2: se ainda restarem placeholders quebrados entre runs, faz merge
      do parágrafo inteiro no run[0] (perda de formatação mínima e localizada).
    """
    doc = Document(modelo_path)

    def substituir_runs(paragraph):
        # Passo 1: substituição dentro de cada run individualmente
        for run in paragraph.runs:
            for chave, valor in dados.items():
                placeholder = f"{{{{{chave}}}}}"
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, str(valor))

        # Passo 2: lida com placeholders quebrados entre múltiplos runs
        full_text = "".join(run.text for run in paragraph.runs)
        if not re.search(r"\{\{[^}]+\}\}", full_text):
            return

        modified = False
        for chave, valor in dados.items():
            placeholder = f"{{{{{chave}}}}}"
            if placeholder in full_text:
                full_text = full_text.replace(placeholder, str(valor))
                modified = True

        if modified and paragraph.runs:
            paragraph.runs[0].text = full_text
            for run in paragraph.runs[1:]:
                run.text = ""

    for paragraph in doc.paragraphs:
        substituir_runs(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    substituir_runs(paragraph)

    doc.save(output_path)


def extrair_placeholders_docx(modelo_path: str) -> list[str]:
    """
    Retorna lista única de placeholders encontrados no .docx, sem chaves.
    Ex.: {{NOME}} => "NOME"
    """
    doc = Document(modelo_path)
    encontrados = set()
    padrao = re.compile(r"\{\{\s*([^{}\s]+)\s*\}\}")

    def coletar_texto(paragraph):
        texto = "".join(run.text for run in paragraph.runs)
        for match in padrao.findall(texto):
            encontrados.add(match.strip())

    for paragraph in doc.paragraphs:
        coletar_texto(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    coletar_texto(paragraph)

    return sorted(encontrados)


def validar_consistencia_modelo(modelo) -> dict:
    """
    Compara os placeholders no .docx com os CampoTemplate ativos.
    Retorna dict com inconsistências para exibir ao administrador.
    """
    try:
        placeholders_docx = set(extrair_placeholders_docx(modelo.arquivo_docx.path))
    except Exception as exc:
        logger.error("Erro ao ler placeholders do docx modelo_id=%s: %s", modelo.pk, exc)
        return {
            "ok": False,
            "erro": str(exc),
            "no_docx_sem_campo": [],
            "no_campo_sem_docx": [],
        }

    campos_ativos = set(
        modelo.campos.filter(ativo=True).values_list("placeholder", flat=True)
    )
    no_docx_sem_campo = sorted(placeholders_docx - campos_ativos)
    no_campo_sem_docx = sorted(campos_ativos - placeholders_docx)

    return {
        "ok": not no_docx_sem_campo and not no_campo_sem_docx,
        "no_docx_sem_campo": no_docx_sem_campo,
        "no_campo_sem_docx": no_campo_sem_docx,
    }
