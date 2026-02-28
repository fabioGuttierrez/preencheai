"""
Script auxiliar para criar um modelo .docx de exemplo com os placeholders corretos.
Execute: python criar_modelo_exemplo.py
O arquivo 'modelo_exemplo.docx' será criado na pasta atual.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def criar_modelo():
    doc = Document()

    # Título
    titulo = doc.add_heading("CONTRATO DE PRESTAÇÃO DE SERVIÇOS", 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    # Identificação das partes
    doc.add_heading("IDENTIFICAÇÃO DAS PARTES", level=1)

    p = doc.add_paragraph()
    p.add_run("CONTRATANTE: ").bold = True
    p.add_run("{{NOME}}")

    p = doc.add_paragraph()
    p.add_run("CPF: ").bold = True
    p.add_run("{{CPF}}")

    p = doc.add_paragraph()
    p.add_run("E-mail: ").bold = True
    p.add_run("{{EMAIL}}")

    p = doc.add_paragraph()
    p.add_run("Telefone: ").bold = True
    p.add_run("{{TELEFONE}}")

    doc.add_paragraph("")

    # Objeto do contrato
    doc.add_heading("CLÁUSULA 1 - OBJETO DO CONTRATO", level=1)

    p = doc.add_paragraph(
        "O presente contrato tem como objeto a prestação de serviços de "
        "{{TIPO_EVENTO}}, a ser realizado em:"
    )

    p = doc.add_paragraph()
    p.add_run("Data: ").bold = True
    p.add_run("{{DATA_EVENTO}}")

    p = doc.add_paragraph()
    p.add_run("Local: ").bold = True
    p.add_run("{{LOCAL_EVENTO}}")

    doc.add_paragraph("")

    # Valor
    doc.add_heading("CLÁUSULA 2 - VALOR E FORMA DE PAGAMENTO", level=1)

    p = doc.add_paragraph()
    p.add_run("O valor total dos serviços é de R$ ").bold = False
    p.add_run("{{VALOR_TOTAL}}").bold = True
    p.add_run(".")

    doc.add_paragraph("")

    # Observações
    doc.add_heading("CLÁUSULA 3 - DISPOSIÇÕES GERAIS", level=1)
    doc.add_paragraph("{{OBSERVACOES}}")

    doc.add_paragraph("")
    doc.add_paragraph("")

    # Assinaturas
    doc.add_paragraph("_" * 40 + "          " + "_" * 40)

    p = doc.add_paragraph()
    p.add_run("CONTRATANTE: {{NOME}}").bold = False
    p.add_run("          ")
    p.add_run("CONTRATADA")

    doc.add_paragraph("")
    doc.add_paragraph(f"Data: ___/___/______")

    doc.save("modelo_exemplo.docx")
    print("✅ Arquivo 'modelo_exemplo.docx' criado com sucesso!")
    print("")
    print("Placeholders disponíveis:")
    placeholders = [
        "{{NOME}}", "{{CPF}}", "{{EMAIL}}", "{{TELEFONE}}",
        "{{TIPO_EVENTO}}", "{{DATA_EVENTO}}", "{{LOCAL_EVENTO}}",
        "{{VALOR_TOTAL}}", "{{OBSERVACOES}}"
    ]
    for p in placeholders:
        print(f"  {p}")


if __name__ == "__main__":
    criar_modelo()
