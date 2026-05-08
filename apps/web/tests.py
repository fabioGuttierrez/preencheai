import os
import tempfile
from unittest.mock import patch

from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings
from django.urls import reverse
from docx import Document

from apps.clientes.models import Cliente
from apps.contratos.models import CampoTemplate, Contrato, ModeloContrato
from apps.core.models import Organizacao, Usuario
from apps.formularios.models import LinkFormulario


def _criar_docx_bytes(texto="Contrato {{VALOR_TOTAL}}"):
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc = Document()
        doc.add_paragraph(texto)
        doc.save(tmp.name)
        tmp_name = tmp.name
    with open(tmp_name, "rb") as f:
        conteudo = f.read()
    os.unlink(tmp_name)
    return conteudo


@override_settings(EMAIL_BACKEND="django.core.mail.backends.locmem.EmailBackend")
class FluxoContratoTest(TestCase):
    def setUp(self):
        self.org = Organizacao.objects.create(nome="Org Teste")
        self.user = Usuario.objects.create_user(
            username="admin",
            email="admin@teste.com",
            password="12345678",
            organizacao=self.org,
        )
        self.cliente = Cliente.objects.create(
            organizacao=self.org,
            nome="Cliente Teste",
            telefone="11999999999",
        )
        self.modelo = ModeloContrato.objects.create(
            organizacao=self.org,
            nome="Modelo Teste",
            arquivo_docx=SimpleUploadedFile(
                "modelo.docx",
                _criar_docx_bytes(),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
        )
        CampoTemplate.objects.create(
            modelo=self.modelo,
            placeholder="VALOR_TOTAL",
            label="Valor Total",
            tipo=CampoTemplate.TIPO_CURRENCY,
            obrigatorio=True,
            ordem=1,
        )
        self.link = LinkFormulario.objects.create(
            organizacao=self.org,
            cliente=self.cliente,
            modelo=self.modelo,
        )

    @patch("services.pdf_service.converter_docx_para_pdf", return_value=None)
    def test_formulario_publico_gera_contrato(self, _mock_pdf):
        url = reverse("formulario_publico", kwargs={"token": self.link.token})
        response = self.client.post(url, {"VALOR_TOTAL": "1500", "aceite_lgpd": "1"})
        self.assertEqual(response.status_code, 200)
        self.assertEqual(Contrato.objects.count(), 1)

    def test_formulario_publico_sem_consentimento(self):
        url = reverse("formulario_publico", kwargs={"token": self.link.token})
        response = self.client.post(url, {"VALOR_TOTAL": "1500"})
        self.assertContains(response, "aceitar os Termos", status_code=200)

    @patch("services.pdf_service.converter_docx_para_pdf", return_value=None)
    def test_placeholder_substituido_no_docx_gerado(self, _mock_pdf):
        """Garante que nenhum {{placeholder}} sobra no arquivo gerado."""
        url = reverse("formulario_publico", kwargs={"token": self.link.token})
        self.client.post(url, {"VALOR_TOTAL": "2500", "aceite_lgpd": "1"})

        contrato = Contrato.objects.get()
        doc = Document(contrato.arquivo_docx.path)
        texto_completo = " ".join(
            "".join(r.text for r in p.runs) for p in doc.paragraphs
        )
        self.assertNotIn("{{", texto_completo, "Placeholder nao substituido encontrado no contrato gerado")
        self.assertIn("2.500", texto_completo)

    @patch("services.pdf_service.converter_docx_para_pdf", return_value=None)
    def test_double_submit_cria_apenas_um_contrato(self, _mock_pdf):
        """Race condition: dois POSTs no mesmo link devem gerar apenas um contrato."""
        url = reverse("formulario_publico", kwargs={"token": self.link.token})
        dados = {"VALOR_TOTAL": "1500", "aceite_lgpd": "1"}

        # Primeiro POST
        r1 = self.client.post(url, dados)
        self.assertEqual(r1.status_code, 200)

        # Segundo POST (link já utilizado)
        r2 = self.client.post(url, dados)
        self.assertEqual(r2.status_code, 200)

        self.assertEqual(Contrato.objects.count(), 1, "Segundo submit nao deve criar contrato duplicado")

    @patch("services.pdf_service.converter_docx_para_pdf", return_value=None)
    def test_link_marcado_como_utilizado(self, _mock_pdf):
        url = reverse("formulario_publico", kwargs={"token": self.link.token})
        self.client.post(url, {"VALOR_TOTAL": "1000", "aceite_lgpd": "1"})
        self.link.refresh_from_db()
        self.assertTrue(self.link.utilizado)

    def test_template_com_arquivo_ausente_retorna_erro_amigavel(self):
        """Se o arquivo .docx do modelo não existe, o usuário vê mensagem clara."""
        self.modelo.arquivo_docx.name = "modelos/arquivo_inexistente.docx"
        self.modelo.save()

        url = reverse("formulario_publico", kwargs={"token": self.link.token})
        response = self.client.post(url, {"VALOR_TOTAL": "1000", "aceite_lgpd": "1"})
        self.assertEqual(response.status_code, 200)
        self.assertEqual(Contrato.objects.count(), 0)


@override_settings(EMAIL_BACKEND="django.core.mail.backends.locmem.EmailBackend")
class ValidacaoConsistenciaTest(TestCase):
    def setUp(self):
        self.org = Organizacao.objects.create(nome="Org Validacao")
        self.user = Usuario.objects.create_user(
            username="admin2", email="a@b.com", password="12345678", organizacao=self.org
        )
        self.client.login(username="admin2", password="12345678")

    def test_upload_modelo_detecta_placeholders(self):
        conteudo = _criar_docx_bytes("Olá {{NOME}}, valor {{VALOR}}")
        url = reverse("modelo_upload")
        response = self.client.post(url, {
            "nome": "Modelo Detecção",
            "arquivo_docx": SimpleUploadedFile(
                "modelo.docx", conteudo,
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
        })
        # Redireciona para configuração de campos após sucesso
        self.assertEqual(response.status_code, 302)
        modelo = ModeloContrato.objects.get(nome="Modelo Detecção")
        placeholders = list(modelo.campos.values_list("placeholder", flat=True))
        self.assertIn("NOME", placeholders)
        self.assertIn("VALOR", placeholders)
